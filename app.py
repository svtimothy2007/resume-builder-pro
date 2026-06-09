"""
Resume Builder Pro - Flask Backend
Production-ready resume builder with PDF/DOCX export, ATS scoring, and SQLite storage.
"""

import os
import json
import sqlite3
import re
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from io import BytesIO

# ReportLab imports for PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate

# python-docx imports for DOCX generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.config['SECRET_KEY'] = 'resume-builder-pro-secret-key-2024'

# ─── Database Setup ───────────────────────────────────────────────────────────

DB_PATH = os.path.join(os.path.dirname(__file__), 'database', 'resumes.db')

def get_db():
    """Get a database connection."""
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """Initialize the database schema."""
    conn = get_db()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS resumes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            data TEXT NOT NULL,
            template TEXT DEFAULT 'classic',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    """Main application page."""
    return render_template('index.html')


@app.route('/save', methods=['POST'])
def save_resume():
    """Save or update a resume in the database."""
    try:
        payload = request.get_json()
        if not payload:
            return jsonify({'error': 'No data provided'}), 400

        name = payload.get('name', 'Untitled Resume')
        data = json.dumps(payload.get('data', {}))
        template = payload.get('template', 'classic')
        resume_id = payload.get('id')
        now = datetime.utcnow().isoformat()

        conn = get_db()
        if resume_id:
            # Update existing
            conn.execute(
                'UPDATE resumes SET name=?, data=?, template=?, updated_at=? WHERE id=?',
                (name, data, template, now, resume_id)
            )
            conn.commit()
            conn.close()
            return jsonify({'id': resume_id, 'message': 'Resume updated successfully'})
        else:
            # Insert new
            cur = conn.execute(
                'INSERT INTO resumes (name, data, template, created_at, updated_at) VALUES (?, ?, ?, ?, ?)',
                (name, data, template, now, now)
            )
            new_id = cur.lastrowid
            conn.commit()
            conn.close()
            return jsonify({'id': new_id, 'message': 'Resume saved successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/resumes', methods=['GET'])
def list_resumes():
    """List all saved resumes."""
    try:
        conn = get_db()
        rows = conn.execute(
            'SELECT id, name, template, created_at, updated_at FROM resumes ORDER BY updated_at DESC'
        ).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/load/<int:resume_id>', methods=['GET'])
def load_resume(resume_id):
    """Load a specific resume by ID."""
    try:
        conn = get_db()
        row = conn.execute('SELECT * FROM resumes WHERE id=?', (resume_id,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Resume not found'}), 404
        result = dict(row)
        result['data'] = json.loads(result['data'])
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/delete/<int:resume_id>', methods=['DELETE'])
def delete_resume(resume_id):
    """Delete a resume by ID."""
    try:
        conn = get_db()
        conn.execute('DELETE FROM resumes WHERE id=?', (resume_id,))
        conn.commit()
        conn.close()
        return jsonify({'message': 'Resume deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/ats-score', methods=['POST'])
def ats_score():
    """Calculate ATS (Applicant Tracking System) score for a resume."""
    try:
        payload = request.get_json()
        data = payload.get('data', {})
        score, breakdown = calculate_ats_score(data)
        return jsonify({'score': score, 'breakdown': breakdown})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/job-match', methods=['POST'])
def job_match():
    """Analyze job description match against resume."""
    try:
        payload = request.get_json()
        data = payload.get('data', {})
        job_description = payload.get('job_description', '')
        result = analyze_job_match(data, job_description)
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/export-pdf', methods=['POST'])
def export_pdf():
    """Export resume as a pixel-perfect PDF."""
    try:
        payload = request.get_json()
        data = payload.get('data', {})
        template = payload.get('template', 'classic')
        pdf_bytes = generate_pdf(data, template)
        name = data.get('personal', {}).get('fullName', 'Resume').replace(' ', '_')
        return send_file(
            BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{name}_Resume.pdf'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/export-docx', methods=['POST'])
def export_docx():
    """Export resume as a DOCX file."""
    try:
        payload = request.get_json()
        data = payload.get('data', {})
        template = payload.get('template', 'classic')
        docx_bytes = generate_docx(data, template)
        name = data.get('personal', {}).get('fullName', 'Resume').replace(' ', '_')
        return send_file(
            BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{name}_Resume.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─── ATS Score Calculator ─────────────────────────────────────────────────────

def calculate_ats_score(data):
    """Calculate comprehensive ATS score with detailed breakdown."""
    score = 0
    breakdown = []

    personal = data.get('personal', {})
    summary = data.get('summary', '')
    skills = data.get('skills', [])
    experience = data.get('experience', [])
    education = data.get('education', [])
    projects = data.get('projects', [])
    certifications = data.get('certifications', [])

    # Contact Information (15 points)
    contact_score = 0
    contact_details = []
    if personal.get('fullName'): contact_score += 3; contact_details.append('Full name present')
    if personal.get('email'): contact_score += 3; contact_details.append('Email present')
    if personal.get('phone'): contact_score += 3; contact_details.append('Phone present')
    if personal.get('linkedin'): contact_score += 3; contact_details.append('LinkedIn present')
    if personal.get('location'): contact_score += 3; contact_details.append('Location present')
    score += contact_score
    breakdown.append({
        'category': 'Contact Information',
        'score': contact_score,
        'max': 15,
        'details': contact_details,
        'icon': '📇'
    })

    # Professional Summary (15 points)
    summary_score = 0
    summary_details = []
    if summary:
        word_count = len(summary.split())
        if word_count >= 50:
            summary_score += 8
            summary_details.append(f'Good length ({word_count} words)')
        elif word_count >= 20:
            summary_score += 5
            summary_details.append(f'Adequate length ({word_count} words)')
        else:
            summary_details.append(f'Too short ({word_count} words, aim for 50+)')
        # Check for action verbs
        action_verbs = ['led', 'managed', 'developed', 'created', 'improved', 'increased',
                       'achieved', 'delivered', 'built', 'designed', 'implemented', 'spearheaded']
        found_verbs = [v for v in action_verbs if v in summary.lower()]
        if found_verbs:
            summary_score += 4
            summary_details.append(f'Contains action verbs: {", ".join(found_verbs[:3])}')
        if any(c.isdigit() for c in summary):
            summary_score += 3
            summary_details.append('Contains quantifiable metrics')
    else:
        summary_details.append('Missing professional summary')
    score += summary_score
    breakdown.append({
        'category': 'Professional Summary',
        'score': summary_score,
        'max': 15,
        'details': summary_details,
        'icon': '📝'
    })

    # Skills Section (20 points)
    skills_score = 0
    skills_details = []
    if skills:
        skill_count = len(skills)
        if skill_count >= 10:
            skills_score += 10
            skills_details.append(f'Excellent skill count ({skill_count} skills)')
        elif skill_count >= 5:
            skills_score += 7
            skills_details.append(f'Good skill count ({skill_count} skills)')
        else:
            skills_score += 3
            skills_details.append(f'Add more skills ({skill_count} skills, aim for 10+)')

        # Check for technical keywords
        tech_keywords = ['python', 'java', 'javascript', 'react', 'sql', 'aws', 'docker',
                        'kubernetes', 'git', 'agile', 'machine learning', 'data analysis']
        found_tech = [k for k in tech_keywords if any(k.lower() in s.lower() for s in skills)]
        if found_tech:
            skills_score += 10
            skills_details.append(f'Technical keywords detected: {", ".join(found_tech[:4])}')
        else:
            skills_details.append('Consider adding technical/industry keywords')
    else:
        skills_details.append('Missing skills section')
    score += skills_score
    breakdown.append({
        'category': 'Skills & Keywords',
        'score': skills_score,
        'max': 20,
        'details': skills_details,
        'icon': '⚡'
    })

    # Work Experience (25 points)
    exp_score = 0
    exp_details = []
    if experience:
        exp_count = len(experience)
        if exp_count >= 3:
            exp_score += 10
            exp_details.append(f'{exp_count} work experiences listed')
        elif exp_count >= 1:
            exp_score += 7
            exp_details.append(f'{exp_count} work experience listed')
        # Check for bullet points / descriptions
        has_descriptions = sum(1 for e in experience if e.get('description') and len(e.get('description', '')) > 30)
        if has_descriptions == exp_count:
            exp_score += 8
            exp_details.append('All positions have detailed descriptions')
        elif has_descriptions > 0:
            exp_score += 4
            exp_details.append(f'{has_descriptions}/{exp_count} positions have descriptions')
        # Check for dates
        has_dates = sum(1 for e in experience if e.get('startDate'))
        if has_dates == exp_count:
            exp_score += 7
            exp_details.append('All positions have dates')
    else:
        exp_details.append('Missing work experience')
    score += exp_score
    breakdown.append({
        'category': 'Work Experience',
        'score': exp_score,
        'max': 25,
        'details': exp_details,
        'icon': '💼'
    })

    # Education (15 points)
    edu_score = 0
    edu_details = []
    if education:
        edu_score += 8
        edu_details.append(f'{len(education)} education entry/entries')
        has_degree = any(e.get('degree') for e in education)
        if has_degree:
            edu_score += 4
            edu_details.append('Degree information present')
        has_dates = any(e.get('graduationDate') or e.get('startDate') for e in education)
        if has_dates:
            edu_score += 3
            edu_details.append('Education dates present')
    else:
        edu_details.append('Missing education section')
    score += edu_score
    breakdown.append({
        'category': 'Education',
        'score': edu_score,
        'max': 15,
        'details': edu_details,
        'icon': '🎓'
    })

    # Bonus Sections (10 points)
    bonus_score = 0
    bonus_details = []
    if projects and len(projects) > 0:
        bonus_score += 3
        bonus_details.append(f'{len(projects)} project(s) listed')
    if certifications and len(certifications) > 0:
        bonus_score += 3
        bonus_details.append(f'{len(certifications)} certification(s) listed')
    if personal.get('github'):
        bonus_score += 2
        bonus_details.append('GitHub profile linked')
    if personal.get('portfolio'):
        bonus_score += 2
        bonus_details.append('Portfolio linked')
    if not bonus_details:
        bonus_details.append('Add projects, certifications, GitHub to boost score')
    score += bonus_score
    breakdown.append({
        'category': 'Bonus Sections',
        'score': bonus_score,
        'max': 10,
        'details': bonus_details,
        'icon': '🏆'
    })

    return min(score, 100), breakdown


# ─── Job Match Analyzer ───────────────────────────────────────────────────────

def analyze_job_match(data, job_description):
    """Analyze how well the resume matches a job description."""
    if not job_description:
        return {'error': 'No job description provided'}

    # Extract all text from resume
    resume_text = extract_resume_text(data)

    # Tokenize job description keywords
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                  'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
                  'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                  'should', 'may', 'might', 'must', 'shall', 'can', 'need', 'you', 'we',
                  'our', 'your', 'their', 'its', 'this', 'that', 'these', 'those'}

    jd_words = re.findall(r'\b[a-zA-Z][a-zA-Z+#.]{1,}\b', job_description.lower())
    jd_keywords = [w for w in jd_words if w not in stop_words and len(w) > 2]
    # Get unique keywords with frequency
    keyword_freq = {}
    for w in jd_keywords:
        keyword_freq[w] = keyword_freq.get(w, 0) + 1

    # Sort by frequency
    sorted_keywords = sorted(keyword_freq.items(), key=lambda x: x[1], reverse=True)
    top_keywords = [k for k, v in sorted_keywords[:30]]

    # Check which keywords are in resume
    resume_lower = resume_text.lower()
    matched = [k for k in top_keywords if k in resume_lower]
    missing = [k for k in top_keywords if k not in resume_lower]

    match_pct = int(len(matched) / max(len(top_keywords), 1) * 100)

    # Categorize missing keywords
    tech_terms = {'python', 'java', 'javascript', 'react', 'node', 'sql', 'aws', 'docker',
                  'kubernetes', 'git', 'api', 'rest', 'graphql', 'typescript', 'angular',
                  'vue', 'django', 'flask', 'spring', 'mongodb', 'postgresql', 'redis',
                  'elasticsearch', 'kafka', 'terraform', 'ci/cd', 'devops', 'agile', 'scrum'}
    missing_tech = [k for k in missing if k in tech_terms]
    missing_soft = [k for k in missing if k not in tech_terms][:10]

    return {
        'match_percentage': match_pct,
        'matched_keywords': matched[:20],
        'missing_keywords': missing[:15],
        'missing_technical': missing_tech[:8],
        'missing_soft_skills': missing_soft[:8],
        'total_jd_keywords': len(top_keywords),
        'recommendation': get_match_recommendation(match_pct)
    }

def extract_resume_text(data):
    """Extract all text from resume data for analysis."""
    parts = []
    personal = data.get('personal', {})
    for v in personal.values():
        if isinstance(v, str): parts.append(v)
    parts.append(data.get('summary', ''))
    for skill in data.get('skills', []):
        parts.append(skill if isinstance(skill, str) else skill.get('name', ''))
    for exp in data.get('experience', []):
        parts.extend([exp.get('company', ''), exp.get('title', ''), exp.get('description', '')])
    for edu in data.get('education', []):
        parts.extend([edu.get('institution', ''), edu.get('degree', ''), edu.get('field', '')])
    for proj in data.get('projects', []):
        parts.extend([proj.get('name', ''), proj.get('description', ''), proj.get('technologies', '')])
    for cert in data.get('certifications', []):
        parts.extend([cert.get('name', ''), cert.get('issuer', '')])
    return ' '.join(parts)

def get_match_recommendation(pct):
    if pct >= 80: return 'Excellent match! Your resume aligns very well with this job.'
    if pct >= 60: return 'Good match. Consider adding a few missing keywords to strengthen your application.'
    if pct >= 40: return 'Moderate match. Tailor your resume by incorporating more job-specific keywords.'
    return 'Low match. Significantly revise your resume to better align with this job description.'


# ─── PDF Generator ────────────────────────────────────────────────────────────

def generate_pdf(data, template='classic'):
    """Generate a pixel-perfect PDF resume using ReportLab."""
    buffer = BytesIO()

    # Template color schemes
    templates = {
        'classic':      {'primary': colors.HexColor('#1a1a2e'), 'accent': colors.HexColor('#e94560'), 'secondary': colors.HexColor('#16213e')},
        'modern':       {'primary': colors.HexColor('#0f3460'), 'accent': colors.HexColor('#e94560'), 'secondary': colors.HexColor('#533483')},
        'minimal':      {'primary': colors.HexColor('#2d3436'), 'accent': colors.HexColor('#0984e3'), 'secondary': colors.HexColor('#636e72')},
        'executive':    {'primary': colors.HexColor('#1b1b2f'), 'accent': colors.HexColor('#c4a35a'), 'secondary': colors.HexColor('#2e2e4e')},
        'creative':     {'primary': colors.HexColor('#6c5ce7'), 'accent': colors.HexColor('#fd79a8'), 'secondary': colors.HexColor('#a29bfe')},
        'tech':         {'primary': colors.HexColor('#00b4d8'), 'accent': colors.HexColor('#90e0ef'), 'secondary': colors.HexColor('#0077b6')},
        'elegant':      {'primary': colors.HexColor('#2c3e50'), 'accent': colors.HexColor('#bdc3c7'), 'secondary': colors.HexColor('#34495e')},
        'bold':         {'primary': colors.HexColor('#d63031'), 'accent': colors.HexColor('#2d3436'), 'secondary': colors.HexColor('#b2bec3')},
        'nature':       {'primary': colors.HexColor('#00b894'), 'accent': colors.HexColor('#55efc4'), 'secondary': colors.HexColor('#00cec9')},
        'professional': {'primary': colors.HexColor('#2b3a55'), 'accent': colors.HexColor('#4a90d9'), 'secondary': colors.HexColor('#3d5a80')},
    }
    colors_map = templates.get(template, templates['classic'])
    primary = colors_map['primary']
    accent = colors_map['accent']
    secondary = colors_map['secondary']

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.6*inch,
        rightMargin=0.6*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Personal Info Header ──
    personal = data.get('personal', {})
    full_name = personal.get('fullName', 'Your Name')
    title = personal.get('title', '')
    email = personal.get('email', '')
    phone = personal.get('phone', '')
    linkedin = personal.get('linkedin', '')
    github = personal.get('github', '')
    portfolio = personal.get('portfolio', '')
    location = personal.get('location', '')

    # Name style
    name_style = ParagraphStyle(
        'Name', fontSize=26, fontName='Helvetica-Bold',
        textColor=primary, spaceAfter=2, leading=30, alignment=TA_CENTER
    )
    title_style = ParagraphStyle(
        'Title', fontSize=12, fontName='Helvetica',
        textColor=accent, spaceAfter=4, leading=16, alignment=TA_CENTER
    )
    contact_style = ParagraphStyle(
        'Contact', fontSize=8.5, fontName='Helvetica',
        textColor=colors.HexColor('#555555'), spaceAfter=0, alignment=TA_CENTER, leading=14
    )

    story.append(Paragraph(full_name, name_style))
    if title:
        story.append(Paragraph(title, title_style))

    # Contact line
    contact_parts = []
    if email: contact_parts.append(f'✉ {email}')
    if phone: contact_parts.append(f'✆ {phone}')
    if location: contact_parts.append(f'⌖ {location}')
    if linkedin: contact_parts.append(f'in {linkedin}')
    if github: contact_parts.append(f'⌨ {github}')
    if portfolio: contact_parts.append(f'⊕ {portfolio}')
    if contact_parts:
        story.append(Paragraph('  |  '.join(contact_parts), contact_style))

    story.append(HRFlowable(width='100%', thickness=2, color=accent, spaceAfter=8, spaceBefore=6))

    # Section header style
    section_header_style = ParagraphStyle(
        'SectionHeader', fontSize=11, fontName='Helvetica-Bold',
        textColor=primary, spaceBefore=10, spaceAfter=4, leading=14
    )
    subsection_bold = ParagraphStyle(
        'SubBold', fontSize=10, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#222222'), spaceAfter=1, leading=13
    )
    subsection_normal = ParagraphStyle(
        'SubNormal', fontSize=9.5, fontName='Helvetica',
        textColor=colors.HexColor('#444444'), spaceAfter=1, leading=13
    )
    body_style = ParagraphStyle(
        'Body', fontSize=9, fontName='Helvetica',
        textColor=colors.HexColor('#333333'), spaceAfter=3, leading=13, leftIndent=10
    )
    italic_style = ParagraphStyle(
        'Italic', fontSize=9, fontName='Helvetica-Oblique',
        textColor=colors.HexColor('#666666'), spaceAfter=2, leading=12
    )

    def add_section_header(title_text):
        story.append(Paragraph(title_text.upper(), section_header_style))
        story.append(HRFlowable(width='100%', thickness=0.8, color=accent, spaceAfter=5))

    # ── Summary ──
    summary = data.get('summary', '')
    if summary:
        add_section_header('Professional Summary')
        summary_style = ParagraphStyle(
            'Summary', fontSize=9.5, fontName='Helvetica',
            textColor=colors.HexColor('#333333'), spaceAfter=4, leading=14, alignment=TA_JUSTIFY
        )
        story.append(Paragraph(summary, summary_style))

    # ── Experience ──
    experience = data.get('experience', [])
    if experience:
        add_section_header('Work Experience')
        for exp in experience:
            company = exp.get('company', '')
            job_title = exp.get('title', '')
            start = exp.get('startDate', '')
            end = exp.get('endDate', 'Present') if not exp.get('current') else 'Present'
            desc = exp.get('description', '')
            location_exp = exp.get('location', '')

            date_str = f'{start} – {end}' if start else ''

            # Company + Date row
            t_data = [[
                Paragraph(f'<b>{company}</b>', subsection_bold),
                Paragraph(date_str, ParagraphStyle('DateR', fontSize=9, fontName='Helvetica',
                    textColor=colors.HexColor('#666666'), alignment=TA_RIGHT, leading=13))
            ]]
            t = Table(t_data, colWidths=[4.5*inch, 2.5*inch])
            t.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
            story.append(t)

            if job_title:
                loc_str = f' — {location_exp}' if location_exp else ''
                story.append(Paragraph(f'{job_title}{loc_str}', italic_style))

            if desc:
                # Parse bullet points
                lines = desc.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line: continue
                    if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        line = line.lstrip('•-* ')
                    story.append(Paragraph(f'• {line}', body_style))

            story.append(Spacer(1, 4))

    # ── Education ──
    education = data.get('education', [])
    if education:
        add_section_header('Education')
        for edu in education:
            inst = edu.get('institution', '')
            degree = edu.get('degree', '')
            field = edu.get('field', '')
            grad = edu.get('graduationDate', '')
            gpa = edu.get('gpa', '')
            degree_field = f'{degree}{", " + field if field else ""}'

            t_data = [[
                Paragraph(f'<b>{inst}</b>', subsection_bold),
                Paragraph(grad, ParagraphStyle('DateR', fontSize=9, fontName='Helvetica',
                    textColor=colors.HexColor('#666666'), alignment=TA_RIGHT, leading=13))
            ]]
            t = Table(t_data, colWidths=[4.5*inch, 2.5*inch])
            t.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
            story.append(t)

            if degree_field:
                gpa_str = f' | GPA: {gpa}' if gpa else ''
                story.append(Paragraph(f'{degree_field}{gpa_str}', italic_style))
            story.append(Spacer(1, 4))

    # ── Skills ──
    skills = data.get('skills', [])
    if skills:
        add_section_header('Skills')
        skill_names = []
        for s in skills:
            if isinstance(s, dict):
                skill_names.append(s.get('name', ''))
            else:
                skill_names.append(str(s))
        skill_text = ' • '.join([s for s in skill_names if s])
        story.append(Paragraph(skill_text, body_style))

    # ── Projects ──
    projects = data.get('projects', [])
    if projects:
        add_section_header('Projects')
        for proj in projects:
            name = proj.get('name', '')
            tech = proj.get('technologies', '')
            url = proj.get('url', '')
            desc = proj.get('description', '')

            header = f'<b>{name}</b>'
            if tech: header += f' <font color="#888888">| {tech}</font>'
            story.append(Paragraph(header, subsection_bold))

            if desc:
                for line in desc.split('\n'):
                    line = line.strip().lstrip('•-* ')
                    if line:
                        story.append(Paragraph(f'• {line}', body_style))
            story.append(Spacer(1, 4))

    # ── Certifications ──
    certifications = data.get('certifications', [])
    if certifications:
        add_section_header('Certifications')
        for cert in certifications:
            name = cert.get('name', '')
            issuer = cert.get('issuer', '')
            date = cert.get('date', '')
            cert_line = f'<b>{name}</b>'
            if issuer: cert_line += f' — {issuer}'
            if date: cert_line += f' ({date})'
            story.append(Paragraph(cert_line, body_style))

    # ── Achievements ──
    achievements = data.get('achievements', [])
    if achievements:
        add_section_header('Achievements')
        for ach in achievements:
            text = ach if isinstance(ach, str) else ach.get('description', '')
            if text:
                story.append(Paragraph(f'• {text}', body_style))

    # ── Leadership ──
    leadership = data.get('leadership', [])
    if leadership:
        add_section_header('Leadership')
        for lead in leadership:
            role = lead.get('role', '')
            org = lead.get('organization', '')
            desc = lead.get('description', '')
            story.append(Paragraph(f'<b>{role}</b>{" — " + org if org else ""}', subsection_bold))
            if desc:
                story.append(Paragraph(desc, body_style))

    # ── Languages ──
    languages = data.get('languages', [])
    if languages:
        add_section_header('Languages')
        lang_parts = []
        for lang in languages:
            if isinstance(lang, dict):
                n = lang.get('name', '')
                level = lang.get('level', '')
                lang_parts.append(f'{n} ({level})' if level else n)
            else:
                lang_parts.append(str(lang))
        story.append(Paragraph(' • '.join(lang_parts), body_style))

    # ── Interests ──
    interests = data.get('interests', [])
    if interests:
        add_section_header('Interests')
        interest_str = ' • '.join([i if isinstance(i, str) else i.get('name', '') for i in interests])
        story.append(Paragraph(interest_str, body_style))

    doc.build(story)
    return buffer.getvalue()


# ─── DOCX Generator ───────────────────────────────────────────────────────────

def generate_docx(data, template='classic'):
    """Generate a professional DOCX resume using python-docx."""
    doc = Document()

    # Template color schemes (hex strings for docx)
    template_colors = {
        'classic':      ('1a1a2e', 'e94560'),
        'modern':       ('0f3460', 'e94560'),
        'minimal':      ('2d3436', '0984e3'),
        'executive':    ('1b1b2f', 'c4a35a'),
        'creative':     ('6c5ce7', 'fd79a8'),
        'tech':         ('00b4d8', '90e0ef'),
        'elegant':      ('2c3e50', 'bdc3c7'),
        'bold':         ('d63031', '2d3436'),
        'nature':       ('00b894', '55efc4'),
        'professional': ('2b3a55', '4a90d9'),
    }
    pri_hex, acc_hex = template_colors.get(template, template_colors['classic'])
    primary_rgb = tuple(int(pri_hex[i:i+2], 16) for i in (0, 2, 4))
    accent_rgb = tuple(int(acc_hex[i:i+2], 16) for i in (0, 2, 4))

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    personal = data.get('personal', {})
    full_name = personal.get('fullName', 'Your Name')
    title = personal.get('title', '')
    email = personal.get('email', '')
    phone = personal.get('phone', '')
    linkedin = personal.get('linkedin', '')
    github = personal.get('github', '')
    location = personal.get('location', '')
    portfolio = personal.get('portfolio', '')

    def set_color(run, rgb):
        run.font.color.rgb = RGBColor(*rgb)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        set_color(run, primary_rgb)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), acc_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Name ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(full_name)
    name_run.bold = True
    name_run.font.size = Pt(24)
    set_color(name_run, primary_rgb)

    if title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        t_run = title_p.add_run(title)
        t_run.font.size = Pt(12)
        set_color(t_run, accent_rgb)

    # Contact info
    contact_parts = []
    if email: contact_parts.append(f'✉ {email}')
    if phone: contact_parts.append(f'✆ {phone}')
    if location: contact_parts.append(f'⌖ {location}')
    if linkedin: contact_parts.append(f'in {linkedin}')
    if github: contact_parts.append(f'⌨ {github}')
    if portfolio: contact_parts.append(f'⊕ {portfolio}')

    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        cr = cp.add_run('  |  '.join(contact_parts))
        cr.font.size = Pt(9)
        set_color(cr, (80, 80, 80))

    # ── Summary ──
    summary = data.get('summary', '')
    if summary:
        add_section_header('Professional Summary')
        sp = doc.add_paragraph(summary)
        sp.paragraph_format.space_after = Pt(4)
        for run in sp.runs:
            run.font.size = Pt(10)

    # ── Experience ──
    experience = data.get('experience', [])
    if experience:
        add_section_header('Work Experience')
        for exp in experience:
            company = exp.get('company', '')
            job_title = exp.get('title', '')
            start = exp.get('startDate', '')
            end = 'Present' if exp.get('current') else exp.get('endDate', 'Present')
            desc = exp.get('description', '')
            location_exp = exp.get('location', '')

            # Company + dates in table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.style = doc.styles['Normal Table']
            table.autofit = False
            table.columns[0].width = Cm(13)
            table.columns[1].width = Cm(5)
            cell_l = table.cell(0, 0)
            cell_r = table.cell(0, 1)
            cell_l.paragraphs[0].clear()
            cr_l = cell_l.paragraphs[0].add_run(company)
            cr_l.bold = True
            cr_l.font.size = Pt(10.5)
            set_color(cr_l, primary_rgb)
            cell_r.paragraphs[0].clear()
            cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_str = f'{start} – {end}' if start else ''
            dr = cell_r.paragraphs[0].add_run(date_str)
            dr.font.size = Pt(9)
            set_color(dr, (100, 100, 100))
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border_el = OxmlElement(f'w:{border_name}')
                        border_el.set(qn('w:val'), 'none')
                        tcBorders.append(border_el)
                    tcPr.append(tcBorders)

            if job_title:
                loc_str = f' — {location_exp}' if location_exp else ''
                tp = doc.add_paragraph(f'{job_title}{loc_str}')
                tp.paragraph_format.space_before = Pt(1)
                tp.paragraph_format.space_after = Pt(2)
                for r in tp.runs:
                    r.font.size = Pt(9.5)
                    r.italic = True
                    set_color(r, (80, 80, 80))

            if desc:
                for line in desc.split('\n'):
                    line = line.strip().lstrip('•-* ')
                    if not line: continue
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.left_indent = Cm(0.5)
                    bp.paragraph_format.space_after = Pt(1)
                    br = bp.add_run(line)
                    br.font.size = Pt(9.5)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Education ──
    education = data.get('education', [])
    if education:
        add_section_header('Education')
        for edu in education:
            inst = edu.get('institution', '')
            degree = edu.get('degree', '')
            field = edu.get('field', '')
            grad = edu.get('graduationDate', '')
            gpa = edu.get('gpa', '')

            table = doc.add_table(rows=1, cols=2)
            table.columns[0].width = Cm(13)
            table.columns[1].width = Cm(5)
            cell_l = table.cell(0, 0)
            cell_r = table.cell(0, 1)
            cell_l.paragraphs[0].clear()
            ir = cell_l.paragraphs[0].add_run(inst)
            ir.bold = True; ir.font.size = Pt(10.5)
            set_color(ir, primary_rgb)
            cell_r.paragraphs[0].clear()
            cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            gr = cell_r.paragraphs[0].add_run(grad)
            gr.font.size = Pt(9); set_color(gr, (100, 100, 100))
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for bn in ['top','left','bottom','right','insideH','insideV']:
                        be = OxmlElement(f'w:{bn}'); be.set(qn('w:val'), 'none'); tcBorders.append(be)
                    tcPr.append(tcBorders)

            if degree or field:
                degree_field = f'{degree}{", " + field if field else ""}'
                gpa_str = f' | GPA: {gpa}' if gpa else ''
                dp = doc.add_paragraph(f'{degree_field}{gpa_str}')
                for r in dp.runs: r.italic = True; r.font.size = Pt(9.5); set_color(r, (80,80,80))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Skills ──
    skills = data.get('skills', [])
    if skills:
        add_section_header('Skills')
        skill_names = [s.get('name','') if isinstance(s, dict) else str(s) for s in skills]
        sp = doc.add_paragraph(' • '.join([s for s in skill_names if s]))
        for r in sp.runs: r.font.size = Pt(9.5)

    # ── Projects ──
    projects = data.get('projects', [])
    if projects:
        add_section_header('Projects')
        for proj in projects:
            pp = doc.add_paragraph()
            pr = pp.add_run(proj.get('name', ''))
            pr.bold = True; pr.font.size = Pt(10.5); set_color(pr, primary_rgb)
            tech = proj.get('technologies', '')
            if tech:
                tr = pp.add_run(f'  |  {tech}')
                tr.font.size = Pt(9); set_color(tr, (100,100,100))
            desc = proj.get('description', '')
            if desc:
                for line in desc.split('\n'):
                    line = line.strip().lstrip('•-* ')
                    if line:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.left_indent = Cm(0.5)
                        br = bp.add_run(line); br.font.size = Pt(9.5)

    # ── Certifications ──
    certifications = data.get('certifications', [])
    if certifications:
        add_section_header('Certifications')
        for cert in certifications:
            name = cert.get('name','')
            issuer = cert.get('issuer','')
            date = cert.get('date','')
            cp = doc.add_paragraph()
            cr = cp.add_run(name); cr.bold = True; cr.font.size = Pt(10)
            if issuer:
                ir = cp.add_run(f' — {issuer}'); ir.font.size = Pt(9.5)
            if date:
                dr = cp.add_run(f' ({date})'); dr.font.size = Pt(9.5); set_color(dr, (100,100,100))

    # ── Achievements ──
    achievements = data.get('achievements', [])
    if achievements:
        add_section_header('Achievements')
        for ach in achievements:
            text = ach if isinstance(ach, str) else ach.get('description','')
            if text:
                ap = doc.add_paragraph(style='List Bullet')
                ar = ap.add_run(text); ar.font.size = Pt(9.5)

    # ── Leadership ──
    leadership = data.get('leadership', [])
    if leadership:
        add_section_header('Leadership')
        for lead in leadership:
            lp = doc.add_paragraph()
            lr = lp.add_run(lead.get('role','')); lr.bold = True; lr.font.size = Pt(10.5)
            if lead.get('organization'):
                or_ = lp.add_run(f' — {lead["organization"]}'); or_.font.size = Pt(9.5)
            if lead.get('description'):
                dp = doc.add_paragraph(lead['description'])
                for r in dp.runs: r.font.size = Pt(9.5)

    # ── Languages ──
    languages = data.get('languages', [])
    if languages:
        add_section_header('Languages')
        lang_parts = []
        for lang in languages:
            if isinstance(lang, dict):
                n = lang.get('name',''); lv = lang.get('level','')
                lang_parts.append(f'{n} ({lv})' if lv else n)
            else:
                lang_parts.append(str(lang))
        lp = doc.add_paragraph(' • '.join(lang_parts))
        for r in lp.runs: r.font.size = Pt(9.5)

    # ── Interests ──
    interests = data.get('interests', [])
    if interests:
        add_section_header('Interests')
        istr = ' • '.join([i if isinstance(i, str) else i.get('name','') for i in interests])
        ip = doc.add_paragraph(istr)
        for r in ip.runs: r.font.size = Pt(9.5)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Entry Point ─────────────────────────────────────────────────────────────

if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)
